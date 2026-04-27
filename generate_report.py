"""生成实验报告Word文档。"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

RESULT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'results')
REPORT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '实验报告.docx')


def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def add_paragraph_styled(doc, text, bold=False, size=12, first_line_indent=True):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p


def add_image_centered(doc, image_path, width=Inches(5.5), caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        run_cap = cap.add_run(caption)
        run_cap.font.size = Pt(10)
        run_cap.font.name = '宋体'
        run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_cap.italic = True


def generate_report():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 封面标题 ====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('基于手工搭建三层神经网络的\nEuroSAT遥感图像分类实验报告')
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run('——基于NumPy实现自动微分与反向传播')
    run2.font.size = Pt(14)
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ==================== 1. 实验概述 ====================
    add_heading_styled(doc, '一、实验概述', level=1)

    add_heading_styled(doc, '1.1 实验目的', level=2)
    add_paragraph_styled(doc,
        '本实验旨在手工搭建一个三层全连接神经网络（MLP）分类器，'
        '在EuroSAT遥感图像数据集上进行训练与测试，实现基于卫星图像的土地覆盖分类任务。'
        '实验要求不使用PyTorch、TensorFlow、JAX等支持自动微分的深度学习框架，'
        '仅使用NumPy进行矩阵运算，自主实现前向传播、反向传播、梯度计算及参数更新等核心功能。')

    add_heading_styled(doc, '1.2 实验环境', level=2)
    add_paragraph_styled(doc,
        '操作系统：Windows 10；编程语言：Python 3；'
        '核心依赖库：NumPy（矩阵运算）、Pillow（图像读取）、Matplotlib（可视化）。'
        '硬件环境：CPU运算，无GPU加速。')

    # ==================== 2. 数据集介绍 ====================
    add_heading_styled(doc, '二、数据集介绍与预处理', level=1)

    add_heading_styled(doc, '2.1 EuroSAT数据集', level=2)
    add_paragraph_styled(doc,
        'EuroSAT是一个基于Sentinel-2卫星图像的土地覆盖分类数据集，'
        '包含10个类别共27,000张64×64像素的RGB彩色图像。各类别及其样本数量如下表所示：')

    # 类别表格
    table = doc.add_table(rows=11, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['类别编号', '类别名称', '样本数量']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    class_data = [
        ('0', 'AnnualCrop（一年生作物）', '3,000'),
        ('1', 'Forest（森林）', '3,000'),
        ('2', 'HerbaceousVegetation（草本植被）', '3,000'),
        ('3', 'Highway（高速公路）', '2,500'),
        ('4', 'Industrial（工业区）', '2,500'),
        ('5', 'Pasture（牧场）', '2,000'),
        ('6', 'PermanentCrop（多年生作物）', '2,500'),
        ('7', 'Residential（住宅区）', '3,000'),
        ('8', 'River（河流）', '2,500'),
        ('9', 'SeaLake（海洋湖泊）', '3,000'),
    ]
    for i, (idx, name, count) in enumerate(class_data):
        set_cell_text(table.rows[i+1].cells[0], idx)
        set_cell_text(table.rows[i+1].cells[1], name, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[i+1].cells[2], count)

    doc.add_paragraph()

    add_heading_styled(doc, '2.2 数据预处理', level=2)
    add_paragraph_styled(doc,
        '数据预处理包括以下步骤：'
        '(1) 图像读取：使用Pillow库读取每张64×64×3的RGB图像；'
        '(2) 像素归一化：将像素值从[0, 255]线性映射到[0.0, 1.0]范围；'
        '(3) 向量展平：将每张图像从64×64×3的三维张量展平为12,288维的一维向量，'
        '作为MLP的输入特征；'
        '(4) 数据集划分：按照7:1.5:1.5的比例，将数据集随机划分为训练集（18,900张）、'
        '验证集（4,050张）和测试集（4,050张），使用固定随机种子（seed=42）保证可复现性。')

    # ==================== 3. 模型结构 ====================
    add_heading_styled(doc, '三、模型结构设计', level=1)

    add_heading_styled(doc, '3.1 网络架构', level=2)
    add_paragraph_styled(doc,
        '本实验采用三层全连接神经网络（MLP），网络结构如下：')
    add_paragraph_styled(doc,
        '输入层 → 全连接层1 → 激活函数 → 全连接层2 → 激活函数 → 全连接层3（输出层） → Softmax',
        bold=True)
    add_paragraph_styled(doc,
        '具体维度配置（基于超参数搜索后的最优配置）：'
        '输入层维度为12,288（= 64×64×3）；'
        '第一隐藏层维度为256，使用Tanh激活函数；'
        '第二隐藏层维度为256，使用Tanh激活函数；'
        '输出层维度为10（对应10个类别），使用Softmax归一化输出概率分布。')

    add_heading_styled(doc, '3.2 权重初始化', level=2)
    add_paragraph_styled(doc,
        '所有全连接层的权重矩阵采用He初始化方法，即权重从均值为0、'
        '标准差为sqrt(2/fan_in)的正态分布中随机采样，其中fan_in为输入特征维度。'
        '偏置向量初始化为全零。He初始化有助于在使用ReLU类激活函数时保持前向传播中信号的方差稳定。')

    add_heading_styled(doc, '3.3 激活函数', level=2)
    add_paragraph_styled(doc,
        '本实验实现了三种激活函数供选择：')
    add_paragraph_styled(doc,
        '(1) ReLU：f(x) = max(0, x)，梯度为x>0时为1，否则为0。'
        'ReLU计算简单，能有效缓解梯度消失问题，是深度学习中最常用的激活函数。')
    add_paragraph_styled(doc,
        '(2) Sigmoid：f(x) = 1/(1+exp(-x))，输出值域为(0,1)。'
        '梯度为f(x)·(1-f(x))，在输入绝对值较大时梯度接近零，容易出现梯度消失。')
    add_paragraph_styled(doc,
        '(3) Tanh：f(x) = tanh(x)，输出值域为(-1,1)。'
        '梯度为1-f(x)²，相比Sigmoid，Tanh的输出以0为中心，'
        '有利于下一层的学习。经超参数搜索，本实验最终选用Tanh作为激活函数。')

    add_heading_styled(doc, '3.4 损失函数', level=2)
    add_paragraph_styled(doc,
        '采用Softmax交叉熵损失函数（Cross-Entropy Loss），将Softmax归一化与交叉熵损失合并计算以保证数值稳定性。'
        '具体实现中使用了log-sum-exp技巧：先将logits减去其最大值，再进行指数运算，'
        '避免指数溢出。交叉熵损失的反向传播梯度简洁优雅：∂L/∂z_i = p_i - y_i，'
        '其中p_i为Softmax输出的概率，y_i为one-hot标签。')

    add_heading_styled(doc, '3.5 反向传播实现', level=2)
    add_paragraph_styled(doc,
        '反向传播基于链式法则，从输出层到输入层逐层传递梯度。每个模块（线性层、激活函数）'
        '在前向传播时缓存必要的中间变量（如输入值、激活输出），在反向传播时利用这些缓存高效计算梯度。'
        '线性层的梯度计算公式为：∂L/∂W = X^T · ∂L/∂Y，∂L/∂b = sum(∂L/∂Y)，'
        '∂L/∂X = ∂L/∂Y · W^T。其中X为层的输入，Y为层的输出。')

    # ==================== 4. 训练策略 ====================
    add_heading_styled(doc, '四、训练策略', level=1)

    add_heading_styled(doc, '4.1 优化器', level=2)
    add_paragraph_styled(doc,
        '采用随机梯度下降（SGD）优化器进行参数更新。参数更新公式为：'
        'W ← W - lr · ∂L/∂W，其中lr为学习率。SGD简单高效，配合学习率衰减策略可以在训练后期获得更精细的参数调整。')

    add_heading_styled(doc, '4.2 学习率衰减', level=2)
    add_paragraph_styled(doc,
        '采用阶梯衰减（Step Decay）策略。初始学习率为0.01，每20个epoch将学习率乘以衰减因子0.5。'
        '即学习率变化为：0.01 → 0.005 → 0.0025 → 0.00125。'
        '学习率衰减有助于训练初期快速收敛，训练后期精细调整参数以避免在最优解附近震荡。')

    add_heading_styled(doc, '4.3 L2正则化', level=2)
    add_paragraph_styled(doc,
        '为防止模型过拟合，在损失函数中加入L2正则化项（Weight Decay），正则化强度λ=1×10⁻⁴。'
        '正则化后的总损失为：L_total = L_CE + (λ/2) · Σ||W||²。'
        'L2正则化通过惩罚过大的权重值，鼓励模型学习更平滑、更具泛化性的特征。')

    add_heading_styled(doc, '4.4 模型保存策略', level=2)
    add_paragraph_styled(doc,
        '在每个epoch结束后，计算模型在验证集上的准确率。'
        '当验证集准确率超过历史最佳值时，自动保存当前模型权重（.npz格式）。'
        '训练结束后，将模型权重恢复为验证集上表现最好的版本，避免因过拟合导致性能下降。')

    # ==================== 5. 超参数搜索 ====================
    add_heading_styled(doc, '五、超参数搜索', level=1)
    add_paragraph_styled(doc,
        '为找到最优的超参数组合，本实验采用网格搜索（Grid Search）方法，'
        '对以下超参数空间进行了全面搜索：')

    # 超参数搜索空间表格
    table2 = doc.add_table(rows=5, cols=2, style='Table Grid')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(table2.rows[0].cells[0], '超参数', bold=True)
    set_cell_text(table2.rows[0].cells[1], '搜索范围', bold=True)
    search_data = [
        ('学习率 (lr)', '0.01, 0.05'),
        ('第一隐藏层大小 (hidden1)', '256, 512'),
        ('第二隐藏层大小 (hidden2)', '128, 256'),
        ('激活函数 (activation)', 'ReLU, Tanh'),
    ]
    for i, (param, values) in enumerate(search_data):
        set_cell_text(table2.rows[i+1].cells[0], param, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table2.rows[i+1].cells[1], values)

    doc.add_paragraph()
    add_paragraph_styled(doc,
        '共计16种超参数组合，每种配置训练20个epoch后在验证集上评估性能。'
        '搜索结果如下图所示（红色柱体为最优配置）：')

    add_image_centered(doc,
        os.path.join(RESULT_DIR, 'search_results.png'),
        width=Inches(5.5),
        caption='图1 网格搜索结果（各配置在验证集上的准确率）')

    add_paragraph_styled(doc,
        '搜索结果分析：')
    add_paragraph_styled(doc,
        '(1) 最优配置：lr=0.01, hidden1=256, hidden2=256, activation=Tanh, weight_decay=1×10⁻⁴，'
        '在20个epoch时达到验证集准确率56.9%。')
    add_paragraph_styled(doc,
        '(2) 学习率对比：lr=0.01的配置整体表现优于lr=0.05。较小的学习率训练更稳定，'
        '不易出现梯度爆炸或震荡。配置10（lr=0.05, hidden1=256, hidden2=256, ReLU）'
        '仅达到11.0%的准确率，说明该组合下学习率过大导致训练发散。')
    add_paragraph_styled(doc,
        '(3) 激活函数对比：在相同架构下，Tanh激活函数总体略优于ReLU。'
        '这可能是因为Tanh输出以0为中心，有利于浅层网络的梯度传递。')
    add_paragraph_styled(doc,
        '(4) 隐藏层大小对比：256×256的组合略优于更大或更小的组合。'
        '过大的隐藏层（如512×256）在参数量显著增加的同时，并未带来明显的精度提升，'
        '反而可能加重过拟合风险。')

    # ==================== 6. 训练过程与实验结果 ====================
    add_heading_styled(doc, '六、训练过程与实验结果', level=1)

    add_heading_styled(doc, '6.1 训练曲线', level=2)
    add_paragraph_styled(doc,
        '使用最优超参数配置（lr=0.01, hidden1=256, hidden2=256, Tanh）训练80个epoch，'
        '训练过程中的损失曲线、准确率曲线和学习率变化如下图所示：')

    add_image_centered(doc,
        os.path.join(RESULT_DIR, 'training_curves.png'),
        width=Inches(6.0),
        caption='图2 训练过程曲线（左：Loss曲线；中：Accuracy曲线；右：学习率衰减）')

    add_paragraph_styled(doc,
        '从Loss曲线可以观察到：'
        '(1) 训练集和验证集的Loss均持续下降，从约2.0降至约1.1，表明模型在有效学习；'
        '(2) 训练集Loss略低于验证集Loss，存在轻微的过拟合现象，但差距不大，'
        '说明L2正则化起到了一定的抑制过拟合效果；'
        '(3) 在每次学习率衰减的节点（第20、40、60个epoch），Loss曲线出现明显的下降加速，'
        '证明学习率衰减策略有效。')

    add_paragraph_styled(doc,
        '从Accuracy曲线可以观察到：'
        '(1) 训练集准确率最终稳定在约64.7%，验证集准确率稳定在约61.4%；'
        '(2) 准确率在前25个epoch增长较快，此后增速放缓并逐渐趋于平稳；'
        '(3) 训练集与验证集的准确率差距约3-5个百分点，表明模型泛化能力尚可。')

    add_heading_styled(doc, '6.2 测试集评估', level=2)
    add_paragraph_styled(doc,
        '在独立测试集（4,050张图像）上，模型的总体分类准确率为60.7%。各类别的详细指标如下表所示：')

    # 测试结果表格
    table3 = doc.add_table(rows=11, cols=5, style='Table Grid')
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    test_headers = ['类别', 'Precision', 'Recall', 'F1-Score', '样本数']
    for i, h in enumerate(test_headers):
        set_cell_text(table3.rows[0].cells[i], h, bold=True, size=9)

    test_data = [
        ('AnnualCrop', '0.5479', '0.5694', '0.5585', '432'),
        ('Forest', '0.7691', '0.8618', '0.8128', '456'),
        ('HerbaceousVegetation', '0.4626', '0.5100', '0.4852', '449'),
        ('Highway', '0.4947', '0.2274', '0.3116', '409'),
        ('Industrial', '0.8256', '0.7809', '0.8026', '388'),
        ('Pasture', '0.5810', '0.7075', '0.6380', '294'),
        ('PermanentCrop', '0.4508', '0.4046', '0.4264', '351'),
        ('Residential', '0.5185', '0.7401', '0.6098', '454'),
        ('River', '0.5696', '0.5771', '0.5733', '376'),
        ('SeaLake', '0.8609', '0.6599', '0.7471', '441'),
    ]
    for i, row_data in enumerate(test_data):
        for j, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table3.rows[i+1].cells[j], val, size=9, align=align)

    doc.add_paragraph()
    add_paragraph_styled(doc,
        '从各类别表现可以看出：'
        '(1) Forest（森林，F1=0.81）和Industrial（工业区，F1=0.80）的分类效果最佳，'
        '这是因为森林具有均匀的绿色纹理，工业区有规则的建筑结构和灰色调，视觉特征鲜明。'
        '(2) SeaLake（海洋湖泊）的Precision最高（0.86），说明被预测为SeaLake的样本大多正确，'
        '但Recall较低（0.66），存在部分海洋湖泊被误分类。'
        '(3) Highway（高速公路）的Recall最低（0.23），说明大量高速公路样本被错误分类为其他类别，'
        '这与高速公路在卫星图像中的视觉多样性有关——城区的高速公路与郊区的有显著差异。')

    add_heading_styled(doc, '6.3 混淆矩阵', level=2)
    add_paragraph_styled(doc,
        '下图为归一化后的混淆矩阵热力图，每行代表真实类别，每列代表预测类别，'
        '数值表示该真实类别被预测为各类别的比例：')

    add_image_centered(doc,
        os.path.join(RESULT_DIR, 'confusion_matrix.png'),
        width=Inches(4.5),
        caption='图3 归一化混淆矩阵')

    add_paragraph_styled(doc,
        '从混淆矩阵可以发现以下易混淆类别对：'
        '(1) Highway常被误分为AnnualCrop、HerbaceousVegetation和River——'
        '高速公路周围常有农田或植被，且河流与公路在卫星图像中都呈现线性结构；'
        '(2) PermanentCrop与HerbaceousVegetation、AnnualCrop相互混淆——'
        '这三种地物类别都以植被为主，在颜色和纹理上具有高度相似性；'
        '(3) Residential有部分被误分为HerbaceousVegetation——'
        '住宅区中的绿化带和花园可能使模型产生混淆。')

    # ==================== 7. 权重可视化 ====================
    add_heading_styled(doc, '七、权重可视化与空间模式分析', level=1)
    add_paragraph_styled(doc,
        '将训练好的第一层隐藏层权重矩阵（维度为12,288×256）的每一列恢复为64×64×3的RGB图像，'
        '可以观察每个神经元学习到的空间特征模式。下图展示了前64个神经元的权重可视化结果：')

    add_image_centered(doc,
        os.path.join(RESULT_DIR, 'weight_visualization.png'),
        width=Inches(5.5),
        caption='图4 第一层隐藏层权重可视化（前64个神经元）')

    add_paragraph_styled(doc,
        '观察权重可视化结果，可以发现以下特征模式：')
    add_paragraph_styled(doc,
        '(1) 色彩倾向性：不同神经元的权重呈现出不同的主��色调。部分神经元（如N0、N1、N16等）'
        '呈现偏绿色调，可能对应森林、牧场等植被覆盖类别的特征响应；'
        '部分神经元（如N23、N27等）呈现偏蓝/紫色调，可能与水体（河流、湖泊）相关；'
        '部分神经元（如N20、N44等）呈现偏暗或混合色调，可能对应建筑物或工业区。')
    add_paragraph_styled(doc,
        '(2) 空间纹理：部分神经元的权重呈现出类似噪声的高频纹理模式，'
        '这在全连接网络中是正常现象——由于MLP将图像展平为一维向量，'
        '丢失了空间局部性信息，因此很难学习到像卷积神经网络（CNN）那样清晰的边缘或纹理滤波器。'
        '但仍有部分神经元呈现出中心-边缘差异或大尺度色块分布，'
        '表明网络在一定程度上捕捉到了全局颜色分布特征。')
    add_paragraph_styled(doc,
        '(3) 与CNN的对比：卷积网络的第一层通常能学到清晰的Gabor滤波器、边缘检测器等局部特征，'
        '而MLP由于缺乏局部感受野的归纳偏置，其权重可视化结果更加分散和全局化。'
        '这也解释了MLP在图像分类任务上性能通常不如CNN的根本原因。')

    # ==================== 8. 错例分析 ====================
    add_heading_styled(doc, '八、错例分析', level=1)
    add_paragraph_styled(doc,
        '从测试集的1,592个分类错误样本中，随机选取16个典型错例进行展示和分析：')

    add_image_centered(doc,
        os.path.join(RESULT_DIR, 'error_analysis.png'),
        width=Inches(5.0),
        caption='图5 分类错误样本示例（红色标题显示真实类别与预测类别）')

    add_paragraph_styled(doc,
        '对上述错例进行分析，可以归纳出以下主要错误原因：')

    add_paragraph_styled(doc,
        '(1) 植被类别间的高度相似性：AnnualCrop（一年生作物）、PermanentCrop（多年生作物）和'
        'HerbaceousVegetation（草本植被）三者在卫星图像中均呈现大面积绿色/棕色植被覆盖，'
        '纹理差异微妙。例如图中"True: PermanentCrop, Pred: HerbaceousVegetation"的样本，'
        '两类在颜色和纹理上几乎无法区分。对于人类而言区分不同类型的农作物也需要专业知识和更高分辨率的图像。')

    add_paragraph_styled(doc,
        '(2) 线性结构的混淆：Highway（高速公路）与River（河流）都呈现线性带状结构。'
        '图中"True: River, Pred: Highway"的样本展示了河流的弯曲走向，'
        '与某些蜿蜒的高速公路在视觉上非常相似。此外，高速公路两侧常有树木和绿化带，'
        '使其进一步与周围的植被类别产生混淆。')

    add_paragraph_styled(doc,
        '(3) 混合用地的歧义：部分样本包含多种地物类型。例如"True: Residential, Pred: HerbaceousVegetation"'
        '的样本中，住宅区被大面积绿色植被包围，模型将其主导色调识别为植被。'
        '又如"True: Highway, Pred: AnnualCrop"的样本，高速公路仅占图像一小部分，'
        '周围的农田面积更大，导致模型做出了错误判断。')

    add_paragraph_styled(doc,
        '(4) 水体与深色地物的混淆：SeaLake（海洋湖泊）与Forest（森林）在特定光照条件下'
        '都呈现深色调，导致图中"True: SeaLake, Pred: Forest"的误分类。'
        '这种错误在色调较暗的湖泊图像中尤其常见。')

    add_paragraph_styled(doc,
        '(5) MLP固有局限性：由于MLP将图像展平为一维向量进行处理，'
        '完全丢失了像素间的空间关系和局部特征（如边缘、角点、纹理方向等）。'
        '这使得模型只能依赖全局颜色分布进行分类，对于视觉特征相似的类别区分能力有限。'
        '使用卷积神经网络（CNN）可以通过局部感受野有效提取空间特征，预期可大幅提升分类精度。')

    # ==================== 9. 总结 ====================
    add_heading_styled(doc, '九、实验总结', level=1)

    add_paragraph_styled(doc,
        '本实验成功使用NumPy从零实现了三层MLP分类器，包括前向传播、反向传播、'
        'SGD优化器、学习率衰减、交叉熵损失、L2正则化等核心模块，'
        '并在EuroSAT遥感图像数据集上完成了训练、验证和测试的完整流程。')

    add_paragraph_styled(doc,
        '主要实验结论如下：')
    add_paragraph_styled(doc,
        '(1) 通过网格搜索确定了最优超参数：学习率0.01、隐藏层大小256×256、Tanh激活函数、'
        '权重衰减系数1×10⁻⁴。模型在测试集上达到60.7%的分类准确率。')
    add_paragraph_styled(doc,
        '(2) 模型对Forest（森林）和Industrial（工业区）等特征鲜明的类别分类效果较好（F1>0.80），'
        '但对Highway（高速公路）和PermanentCrop（多年生作物）等视觉特征模糊的类别表现欠佳（F1<0.43）。')
    add_paragraph_styled(doc,
        '(3) 权重可视化结果显示MLP的第一层主要学习了全局色调特征，难以提取精细的空间纹理模式。'
        '这是MLP相比CNN在图像分类任务上的核心劣势。')
    add_paragraph_styled(doc,
        '(4) 错例分析揭示了植被类别间高度相似、线性结构混淆、混合用地歧义等问题，'
        '这些问题在一定程度上可以通过引入卷积操作和数据增强等技术来缓解。')

    add_paragraph_styled(doc,
        '可能的改进方向：'
        '(1) 引入数据增强（随机翻转、旋转、颜色抖动）以扩充训练数据多样性；'
        '(2) 增加隐藏层深度或使用Batch Normalization改善训练稳定性；'
        '(3) 使用PCA或其他降维方法减少输入维度，降低过拟合风险；'
        '(4) 采用更先进的优化器（如Adam）或学习率调度策略（如余弦退火）；'
        '(5) 最根本的改进是使用卷积神经网络（CNN），利用局部感受野和权重共享机制'
        '有效提取空间特征，预期可将准确率提升至85%以上。')

    # 保存文档
    doc.save(REPORT_PATH)
    print(f"实验报告已保存至：{REPORT_PATH}")


if __name__ == '__main__':
    generate_report()
